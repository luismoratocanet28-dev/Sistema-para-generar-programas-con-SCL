import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_word_report(md_file, output_file):
    doc = docx.Document()
    
    # Title
    title = doc.add_heading('Investigación: Estándar IEC 61131-3 y Siemens SCL', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), level=1)
        elif line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), level=2)
        elif line.startswith('*   **'):
            p = doc.add_paragraph(style='List Bullet')
            # Extract bold text if any
            parts = line.split('**')
            for i, part in enumerate(parts):
                if i % 2 == 1: # Bold part
                    p.add_run(part).bold = True
                else:
                    p.add_run(part.replace('* ', '', 1))
        elif line.startswith('```'):
            continue # Simple skipping code blocks for now or handling them
        else:
            doc.add_paragraph(line)
            
    doc.save(output_file)
    print(f"Documento generado en: {output_file}")

if __name__ == "__main__":
    base_path = r'c:\Users\portatil\OneDrive\Documentos\Escritorio\Informatica\Sistema para generar programas con SCL'
    md_input = os.path.join(base_path, 'docs', 'investigacion_iec61131.md')
    docx_output = os.path.join(base_path, 'docs', 'investigacion_iec61131.docx')
    
    create_word_report(md_input, docx_output)
